import os
import json
from sqlalchemy.orm import Session
import models
from services.ai_service import get_embeddings_batch, extract_doc_metadata
from services.storage_service import storage_service

# Standard libraries for various formats
import docx
import openpyxl
from pypdf import PdfReader

def analyze_document(document_id: str, db: Session):
    """Extract text from various formats (PDF, DOCX, XLSX, TXT), chunk it, embed it, and save to DB."""
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        return

    try:
        doc.status = "processing"
        db.commit()

        # In local dev, the file should already be in /app/uploads/ (shared via volume or direct write)
        # Check local path first
        local_upload_path = f"/app/uploads/{doc.id}.{doc.file_name.split('.')[-1].lower()}"
        file_path = f"/tmp/{doc.storage_path}"

        if os.path.exists(local_upload_path):
            print(f"DEBUG: Using local file found at {local_upload_path}")
            file_path = local_upload_path
        elif not os.path.exists(file_path):
            print(f"DEBUG: Downloading {doc.storage_path} from Supabase...")
            os.makedirs("/tmp", exist_ok=True)
            success = storage_service.download_file(doc.storage_path, file_path)
            if not success:
                raise Exception("Failed to download file from storage")

        # ... (rest of extraction logic remains same)
        file_ext = doc.file_name.split(".")[-1].lower()
        chunks = []
        # ...

        print(f"DEBUG: Extracting text from {doc.file_name} (format: {file_ext})...")

        if file_ext == "pdf":
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    chunks.append(text)
        
        elif file_ext == "docx":
            doc_obj = docx.Document(file_path)
            # Group paragraphs into chunks (already implemented below)
            # Group paragraphs into chunks to avoid too many small chunks
            full_text = []
            for para in doc_obj.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Simple chunking by paragraph groups or character count
            current_chunk = ""
            for text in full_text:
                if len(current_chunk) + len(text) > 1500:
                    chunks.append(current_chunk)
                    current_chunk = text
                else:
                    current_chunk += "\n" + text if current_chunk else text
            if current_chunk:
                chunks.append(current_chunk)

        elif file_ext == "xlsx":
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.worksheets:
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty rows
                    if any(cell is not None for cell in row):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        sheet_data.append(row_text)
                
                if sheet_data:
                    # One chunk per sheet or split if too large
                    sheet_text = f"Sheet: {sheet.title}\n" + "\n".join(sheet_data)
                    # For very large sheets, we might need smaller chunks, but for MVP one per sheet is a start
                    if len(sheet_text) > 5000:
                        # Simple split
                        for i in range(0, len(sheet_text), 4000):
                            chunks.append(sheet_text[i:i+4000])
                    else:
                        chunks.append(sheet_text)

        elif file_ext in ["txt", "md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                # Split by roughly 1500 characters with some overlap or line-based
                for i in range(0, len(content), 1500):
                    chunks.append(content[i:i+1500])

        print(f"DEBUG: Extracted {len(chunks)} chunks of text.")
        
        if chunks:
            print(f"DEBUG: Generating embeddings for {len(chunks)} chunks...")
            # Generate embeddings in batches
            batch_size = 100
            for i in range(0, len(chunks), batch_size):
                batch_slice = chunks[i:i + batch_size]
                embeddings = get_embeddings_batch(batch_slice)
                
                if embeddings:
                    print(f"DEBUG: Saving batch {i//batch_size + 1} to database...")
                    for j, (text, emb) in enumerate(zip(batch_slice, embeddings)):
                        chunk = models.DocumentChunk(
                            document_id=doc.id,
                            chunk_index=i + j,
                            content=text,
                            embedding=emb
                        )
                        db.add(chunk)
                        db.add(chunk)
            db.commit()
        
        # 4. Extract Metadata (Summary, Tags, etc.)
        if chunks:
            full_text_sample = "\n".join(chunks)
            print(f"DEBUG: Extracting metadata for {doc.file_name}...")
            metadata = extract_doc_metadata(full_text_sample)
            
            doc.document_type = metadata.get("document_type", "未分類")
            doc.customer_name = metadata.get("customer_name", "")
            # 保存時はJSON形式で、短い概要と詳細レポートの両方を保持
            doc.summary = json.dumps({
                "brief": metadata.get("summary", ""),
                "detailed": metadata.get("content_report", "")
            }, ensure_ascii=False)
            doc.tags = metadata.get("tags", "")
            db.commit()

        doc.status = "completed"
        db.commit()
        print(f"DEBUG: Analysis COMPLETED for {doc.file_name}")
    except Exception as e:
        print(f"Error analyzing document {document_id}: {e}")
        db.rollback()
        raise e

def reextract_document_tags(document_id: str, db: Session):
    """Re-extract metadata (tags, customer_name, summary) from the original file."""
    doc = db.query(models.Document).filter(models.Document.id == document_id).first()
    if not doc:
        return

    try:
        doc.status = "processing"
        db.commit()

        # 1. 元ファイルパスの特定・ダウンロード
        local_upload_path = f"/app/uploads/{doc.id}.{doc.file_name.split('.')[-1].lower()}"
        file_path = f"/tmp/{doc.storage_path}"

        if os.path.exists(local_upload_path):
            print(f"DEBUG: Using local file found at {local_upload_path}")
            file_path = local_upload_path
        elif not os.path.exists(file_path):
            print(f"DEBUG: Downloading {doc.storage_path} from Supabase...")
            os.makedirs("/tmp", exist_ok=True)
            success = storage_service.download_file(doc.storage_path, file_path)
            if not success:
                raise Exception("Failed to download file from storage")

        file_ext = doc.file_name.split(".")[-1].lower()
        chunks = []

        print(f"DEBUG: Re-extracting text from {doc.file_name} (format: {file_ext})...")

        if file_ext == "pdf":
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    chunks.append(text)
        
        elif file_ext == "docx":
            doc_obj = docx.Document(file_path)
            full_text = []
            for para in doc_obj.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            current_chunk = ""
            for text in full_text:
                if len(current_chunk) + len(text) > 1500:
                    chunks.append(current_chunk)
                    current_chunk = text
                else:
                    current_chunk += "\n" + text if current_chunk else text
            if current_chunk:
                chunks.append(current_chunk)

        elif file_ext == "xlsx":
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.worksheets:
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        sheet_data.append(row_text)
                
                if sheet_data:
                    sheet_text = f"Sheet: {sheet.title}\n" + "\n".join(sheet_data)
                    if len(sheet_text) > 5000:
                        for i in range(0, len(sheet_text), 4000):
                            chunks.append(sheet_text[i:i+4000])
                    else:
                        chunks.append(sheet_text)

        elif file_ext in ["txt", "md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                for i in range(0, len(content), 1500):
                    chunks.append(content[i:i+1500])

        if chunks:
            full_text_sample = "\n".join(chunks)
            print(f"DEBUG: Re-extracting metadata for {doc.file_name}...")
            metadata = extract_doc_metadata(full_text_sample)
            
            doc.document_type = metadata.get("document_type", "未分類")
            doc.customer_name = metadata.get("customer_name", "")
            doc.summary = json.dumps({
                "brief": metadata.get("summary", ""),
                "detailed": metadata.get("content_report", "")
            }, ensure_ascii=False)
            doc.tags = metadata.get("tags", "")
            
        doc.status = "completed"
        db.commit()
        print(f"DEBUG: Tag re-extraction COMPLETED for {doc.file_name}")
    except Exception as e:
        print(f"Error re-extracting document tags {document_id}: {e}")
        db.rollback()
        try:
            doc.status = "failed"
            db.commit()
        except:
            pass
        raise e

